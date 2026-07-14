#!/usr/bin/env python3
"""End-to-End-Test der fokussierten Schriftsatz-Versandwerkstatt."""

from __future__ import annotations

import importlib.util
import json
import shutil
import sys
import tempfile
from email.message import EmailMessage
from pathlib import Path
from types import ModuleType

from docx import Document
from openpyxl import Workbook
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


REPO = Path(__file__).resolve().parent.parent
TOOL = (
    REPO
    / "schriftsatz-versandwerkstatt"
    / "skills"
    / "versandmappe-endfertigen"
    / "werkzeuge"
    / "build_versandmappe.py"
)


def load_tool() -> ModuleType:
    spec = importlib.util.spec_from_file_location("build_versandmappe", TOOL)
    if spec is None or spec.loader is None:
        raise AssertionError(f"Werkzeug nicht ladbar: {TOOL}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def write_pdf(path: Path, title: str, pages: int) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    _, height = A4
    for page in range(1, pages + 1):
        pdf.setFont("Helvetica-Bold", 15)
        pdf.drawString(56, height - 72, title)
        pdf.setFont("Helvetica", 11)
        pdf.drawString(56, height - 106, f"Seite {page} von {pages}")
        for row in range(8):
            pdf.drawString(56, height - 145 - row * 22, f"Dokumentinhalt {row + 1} auf Seite {page}.")
        pdf.showPage()
    pdf.save()


def write_eml(path: Path) -> None:
    message = EmailMessage()
    message["From"] = "bauleitung@example.org"
    message["To"] = "kanzlei@example.org"
    message["Date"] = "Tue, 14 Jul 2026 09:30:00 +0200"
    message["Subject"] = "Abnahme und Restarbeiten"
    message.set_content("Die Abnahme fand am 13. Juli 2026 statt. Die Restarbeiten sind bis zum 20. Juli auszuführen.")
    path.write_bytes(message.as_bytes(policy=message.policy))


def write_office_sources(docx_path: Path, xlsx_path: Path) -> None:
    document = Document()
    document.add_heading("Klageerwiderung", level=1)
    document.add_paragraph("Landgericht Essen")
    document.add_paragraph("Aktenzeichen 12 O 34/26")
    document.add_paragraph("Der Beklagte beantragt, die Klage abzuweisen.")
    document.add_paragraph("Rechtsanwalt Max Muster")
    document.save(docx_path)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Berechnung"
    sheet.append(["Position", "Betrag", "Datum"])
    sheet.append(["Hauptforderung", 12500, "2026-06-30"])
    sheet.append(["Zahlung", -2500, "2026-07-03"])
    workbook.save(xlsx_path)


def main() -> int:
    tool = load_tool()
    with tempfile.TemporaryDirectory(prefix="versandwerkstatt-test-") as tmp:
        root = Path(tmp)
        source = root / "eingang"
        target = root / "ausgang"
        source.mkdir()

        office_test = bool(shutil.which("soffice") or shutil.which("libreoffice"))
        if office_test:
            lead = source / "Klageerwiderung.docx"
            write_office_sources(lead, source / "Anlage_B-04_Berechnung.xlsx")
        else:
            lead = source / "Klageerwiderung.pdf"
            write_pdf(lead, "Klageerwiderung", 2)
        write_pdf(source / "Anlage_B-01_Kaufvertrag.pdf", "Kaufvertrag", 2)
        write_pdf(source / "Anlage_B-02_Mahnung.pdf", "Mahnung", 1)
        email_dir = source / "e-mails"
        email_dir.mkdir()
        write_eml(email_dir / "Anlage_B-03_E-Mail-Abnahme.eml")

        exit_code = tool.main(
            [
                "--eingang",
                str(source),
                "--ausgang",
                str(target),
                "--praefix",
                "B",
                "--hauptdokument",
                str(lead),
                "--dokumentart",
                "Klageerwiderung_12_O_34_26",
                "--gericht",
                "Landgericht Essen",
                "--aktenzeichen",
                "12 O 34/26",
                "--frist",
                "2026-07-15 23:59",
                "--verantwortlich",
                "Rechtsanwalt Max Muster",
                "--versender",
                "Max Muster",
                "--signaturweg",
                "persoenlich-sicher",
                "--sichtpruefung-bestaetigt",
                "--strict",
            ]
        )
        if exit_code != 0:
            raise AssertionError(f"Werkzeuglauf endete mit Status {exit_code}")

        shipping = target / "versandfertig"
        shipped = sorted(shipping.glob("*.pdf"))
        expected_files = 5 if office_test else 4
        if len(shipped) != expected_files:
            raise AssertionError(f"{expected_files} Versanddateien erwartet, erhalten: {shipped}")
        if any(len(path.name) > 80 or not path.name.isascii() or " " in path.name for path in shipped):
            raise AssertionError(f"Dateinamensprofil verletzt: {[p.name for p in shipped]}")

        email_pdf = next(path for path in shipped if "AnlageB3" in path.name)
        email_text = "\n".join((page.extract_text() or "") for page in PdfReader(str(email_pdf)).pages)
        for marker in ("Anlage B 3", "Abnahme und Restarbeiten", "bauleitung@example.org"):
            if marker not in email_text:
                raise AssertionError(f"E-Mail-PDF enthält {marker!r} nicht")

        manifest = json.loads((target / "intern" / "Versandmanifest.json").read_text(encoding="utf-8"))
        if manifest["metadaten"]["signaturweg"] != "persoenlich-sicher":
            raise AssertionError("Signaturweg fehlt im Manifest")
        email_entry = next(entry for entry in manifest["anlagen"] if entry["anlage"] == "Anlage B 3")
        if email_entry["quelle"] != "e-mails/Anlage_B-03_E-Mail-Abnahme.eml":
            raise AssertionError("Relativer Quellpfad der E-Mail fehlt im Manifest")
        for name in ("Freigabevermerk.md", "Eingangskontrolle.md", "Preflight-Bericht.md"):
            if not (target / "intern" / name).is_file():
                raise AssertionError(f"Interne Prüfausgabe fehlt: {name}")
        if "| STOP |" in (target / "intern" / "Preflight-Bericht.md").read_text(encoding="utf-8"):
            raise AssertionError("Preflight enthält unerwarteten Stop-Befund")

        mismatch = root / "mismatch"
        mismatch_code = tool.main(
            [
                "--eingang",
                str(source),
                "--ausgang",
                str(mismatch),
                "--praefix",
                "B",
                "--hauptdokument",
                str(lead),
                "--gericht",
                "Landgericht Essen",
                "--aktenzeichen",
                "12 O 34/26",
                "--frist",
                "2026-07-15 23:59",
                "--verantwortlich",
                "Rechtsanwalt Max Muster",
                "--versender",
                "Kanzleimitarbeiterin Jana Winter",
                "--signaturweg",
                "persoenlich-sicher",
                "--sichtpruefung-bestaetigt",
                "--strict",
            ]
        )
        if mismatch_code != 3:
            raise AssertionError("Fremdversand ohne qualifizierte elektronische Signatur wurde nicht gestoppt")
        mismatch_preflight = (mismatch / "intern" / "Preflight-Bericht.md").read_text(encoding="utf-8")
        if "stimmen beim persönlichen sicheren Versand nicht überein" not in mismatch_preflight:
            raise AssertionError("Absenderkonflikt fehlt im Preflight")

        qes = root / "qes"
        qes_code = tool.main(
            [
                "--eingang",
                str(source),
                "--ausgang",
                str(qes),
                "--praefix",
                "B",
                "--hauptdokument",
                str(lead),
                "--gericht",
                "Landgericht Essen",
                "--aktenzeichen",
                "12 O 34/26",
                "--frist",
                "2026-07-15 23:59",
                "--verantwortlich",
                "Rechtsanwalt Max Muster",
                "--versender",
                "Kanzleimitarbeiterin Jana Winter",
                "--signaturweg",
                "qes",
                "--sichtpruefung-bestaetigt",
                "--strict",
            ]
        )
        if qes_code != 3:
            raise AssertionError("Nicht bestätigte qualifizierte elektronische Signatur wurde nicht gestoppt")
        qes_preflight = (qes / "intern" / "Preflight-Bericht.md").read_text(encoding="utf-8")
        if "qualifizierte elektronische Signatur ist nicht als manuell geprüft bestätigt" not in qes_preflight:
            raise AssertionError("Signaturprüf-Stop fehlt im Preflight")

    office_status = "mit Office-Konvertierung" if office_test else "ohne verfügbares LibreOffice"
    print(f"test-schriftsatz-versandwerkstatt OK (Positivlauf {office_status}, EML, Absenderkonflikt und Signatur-Stop)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
