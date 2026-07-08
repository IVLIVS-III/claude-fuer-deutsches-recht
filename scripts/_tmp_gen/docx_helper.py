"""Wiederverwendbares Hilfsmodul zur Erzeugung von Testakten-DOCX-Aktenstuecken
im Stil der bestehenden Akten (Times New Roman 11pt, dezimale Gliederung).

Nicht Teil des Produktiv-Repos - nur ein Bauwerkzeug fuer diese Ausbau-Session.
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def new_document():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    return doc


def add_title(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    return p


def add_h2(doc, text):
    return doc.add_paragraph(text, style="Heading 2")


def add_h3(doc, text):
    return doc.add_paragraph(text, style="Heading 3")


def add_p(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_letterhead(doc, lines):
    """lines: list of strings, first is bold (Kanzlei/Behoerdenname)."""
    for i, line in enumerate(lines):
        add_p(doc, line, bold=(i == 0))
    doc.add_paragraph()


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_formathinweis(doc):
    add_p(doc, "Formathinweis: Dieses Aktenstueck wird in Times New Roman, 11 pt, mit ausschliesslich dezimaler Gliederung ausgegeben.".replace("Aktenstueck", "Aktenstück").replace("ausschliesslich", "ausschließlich"))


def save(doc, path):
    doc.save(path)
    print("geschrieben:", path)
