#!/usr/bin/env python3
"""Wiederverwendbares Toolkit zur Erzeugung realistischer Testakten-Aktenstuecke.

Nur fuer den internen Ausbau der Insolvenzrecht-Testakten. Erzeugt DOCX,
EML, CSV, XLSX und einfache JPG-Belege im Stil der bestehenden Akten.
"""
from __future__ import annotations

import csv
import io
import os
import random
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from PIL import Image, ImageDraw, ImageFont

REPO_ROOT = Path(__file__).resolve().parent.parent.parent
TESTAKTEN = REPO_ROOT / "testakten"


def make_docx(path: Path, kopf: str, titel: str, absaetze: list[str], unterschrift: str | None = None):
    """Erzeugt ein DOCX-Aktenstueck im bestehenden Kanzleistil."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph(kopf)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = None

    doc.add_heading(titel, level=1)

    for a in absaetze:
        if a.startswith("## "):
            doc.add_heading(a[3:], level=2)
        elif a.startswith("- "):
            doc.add_paragraph(a[2:], style="List Bullet")
        else:
            doc.add_paragraph(a)

    if unterschrift:
        doc.add_paragraph("")
        doc.add_paragraph(unterschrift)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def make_eml(path: Path, frm: str, to: str, subject: str, date: str, body: str, msgid_domain: str = "kanzlei-example.de"):
    path.parent.mkdir(parents=True, exist_ok=True)
    msgid = f"<{abs(hash(subject+date))%10**12}-{random.randint(1000,9999)}@{msgid_domain}>"
    content = (
        f"From: {frm}\n"
        f"To: {to}\n"
        f"Date: {date}\n"
        f"Subject: {subject}\n"
        f"Message-ID: {msgid}\n"
        "MIME-Version: 1.0\n"
        "Content-Type: text/plain; charset=UTF-8\n"
        "Content-Transfer-Encoding: quoted-printable\n"
        "\n"
        f"{body}\n"
    )
    path.write_text(content, encoding="utf-8")
    return path


def make_md_email(path: Path, frm: str, to: str, subject: str, date: str, body: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    content = (
        f"**From:** {frm}  \n"
        f"**To:** {to}  \n"
        f"**Date:** {date}  \n"
        f"**Subject:** {subject}\n\n"
        "---\n\n"
        f"{body}\n"
    )
    path.write_text(content, encoding="utf-8")
    return path


def make_csv(path: Path, header: list[str], rows: list[list], delimiter: str = ";"):
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f, delimiter=delimiter)
        w.writerow(header)
        for r in rows:
            w.writerow(r)
    return path


def make_xlsx(path: Path, sheet_name: str, header: list[str], rows: list[list], title: str | None = None):
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31]
    start_row = 1
    if title:
        ws["A1"] = title
        ws["A1"].font = Font(bold=True, size=13)
        start_row = 3
    hdr_fill = PatternFill(start_color="01696F", end_color="01696F", fill_type="solid")
    for ci, h in enumerate(header, start=1):
        c = ws.cell(row=start_row, column=ci, value=h)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = hdr_fill
    for ri, row in enumerate(rows, start=start_row + 1):
        for ci, val in enumerate(row, start=1):
            ws.cell(row=ri, column=ci, value=val)
    for ci, h in enumerate(header, start=1):
        ws.column_dimensions[chr(64 + ci) if ci <= 26 else "A"].width = max(14, len(str(h)) + 4)
    wb.save(path)
    return path


def make_jpg(path: Path, title: str, lines: list[str], size=(1000, 700)):
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", size, color=(245, 244, 240))
    d = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
        font_body = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
    except Exception:
        font_title = ImageFont.load_default()
        font_body = ImageFont.load_default()
    d.rectangle([(0, 0), (size[0], 90)], fill=(1, 105, 111))
    d.text((30, 30), title, fill=(255, 255, 255), font=font_title)
    y = 130
    for line in lines:
        d.text((30, y), line, fill=(40, 40, 40), font=font_body)
        y += 34
    d.rectangle([(0, size[1]-40), (size[0], size[1])], fill=(212, 209, 202))
    d.text((30, size[1]-34), "Anlage zur Akte - interne Dokumentation", fill=(80, 80, 80), font=font_body)
    img.save(path, "JPEG", quality=88)
    return path


if __name__ == "__main__":
    print("aktenbau toolkit geladen")
